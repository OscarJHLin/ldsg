#!/usr/bin/env python3
"""生成 LDSG 论文 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# 标题
title = doc.add_heading('分层动态空间图记忆系统', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Hierarchical Dynamic Spatial Graph Memory System (LDSG)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True
doc.add_paragraph('版本 2.0 | 2026年4月3日')
doc.add_paragraph('协作者：林筠贺、Kimi')
doc.add_page_break()

# 摘要
doc.add_heading('摘要', level=1)
doc.add_paragraph(
    '当前大语言模型（LLM）面临两大根本困境：上下文窗口限制与静态知识表示。'
    '本文提出分层动态空间图记忆系统（LDSG），一种受认知科学启发的统一记忆框架。'
    '该系统通过三层架构（STM工作记忆/MSG主图空间/WS工作空间）+ 三锚点动态压缩机制，'
    '实现持续学习、渐进遗忘与创造性联想的有机融合。'
)

# 1. 引言
doc.add_heading('1. 引言', level=1)
doc.add_heading('1.1 问题背景', level=2)
doc.add_paragraph('LLM 的两个结构性缺陷：')
for item in [
    '上下文窗口瓶颈：Transformer O(n²) 自注意力导致上下文长度受限',
    '参数知识的静态性：知识固化于预训练权重，无法实时适应个体用户',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.2 核心创新', level=2)
for item in [
    '空间-语义统一表示：相关性查询 O(1) 复杂度',
    '关键-子节点分层连接：显式+隐式双模式',
    '三锚点动态压缩栈：多任务并行与思维广度增强',
    '渐进式结构化遗忘：从完整子图到单概念线索的层级退化',
]:
    doc.add_paragraph(item, style='List Bullet')

# 2. 认知科学基础
doc.add_heading('2. 认知科学基础', level=1)
headers = ['概念', '认知科学来源', 'LDSG对应机制']
rows = [
    ['工作记忆容量限制', 'Baddeley (2000)', '三锚点清晰区（max=3）'],
    ['系统巩固', 'Dudai (2004)', '关键-子节点分层归并'],
    ['默认模式网络', 'Raichle (2001)', '三锚点协同漫游（focused/scan/bridge）'],
    ['熟悉感vs回忆', 'Bjork (1992)', '压缩区激活度+唤起成本量化'],
]
table = doc.add_table(rows=1+len(rows), cols=3)
table.style = 'Table Grid'
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for ri, row in enumerate(rows):
    for ci, cell in enumerate(row):
        table.rows[ri+1].cells[ci].text = cell
doc.add_paragraph()

# 3. 三层架构
doc.add_heading('3. 三层架构详解', level=1)

doc.add_heading('3.1 短期记忆层（STM）', level=2)
doc.add_paragraph(
    '子图结构：关键节点（1-3个）携带子节点（5-20个），边包含关系类型、权重与情境向量。'
    '归并过程：(1) 关键节点在MSG中定位或创建；(2) 建立显式边连接；'
    '(3) 子节点压缩为签名附加到宿主关键节点；(4) 触发局部力导向松弛。'
)

doc.add_heading('3.2 主图空间层（MSG）', level=2)
doc.add_paragraph(
    '空间表示：384维单位向量，距离 d(i,j)=|vi-vj|，相关性 sim(i,j)≈1-d²/2。'
)
doc.add_paragraph('分层结构：')
for item in [
    'L3元层（~10²节点）：自我、工作、家庭等最高层概念',
    'L2关键层（~10⁴节点）：具体项目、人物、事件',
    'L1子节点层（~10⁶节点）：细节属性，压缩包形式依附L2',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.3 三锚点动态压缩栈', level=2)
doc.add_paragraph('三锚点权重分配：')
rows = [
    ['primary', '60%', '深度挖掘（depth=3）', '当前核心话题'],
    ['secondary', '25%', '广度扫描（depth=1）', '相关背景/并行任务'],
    ['tertiary', '15%', '边界检测（depth=0）', '潜在关联/待办事项'],
]
table = doc.add_table(rows=1+len(rows), cols=4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '角色'
table.rows[0].cells[1].text = '权重'
table.rows[0].cells[2].text = '漫游模式'
table.rows[0].cells[3].text = '功能'
for ri, row in enumerate(rows):
    for ci, cell in enumerate(row):
        table.rows[ri+1].cells[ci].text = cell
doc.add_paragraph()

doc.add_paragraph('动态压缩栈层级：')
for item in [
    '清晰区（max=3）：完整子图，激活即时可访问',
    '压缩区（max=6）：摘要(summary)→线索(clue)层级，中等唤起成本',
    '遗忘区：激活度低于阈值 → 从栈中清除（MSG节点本身不受影响）',
]:
    doc.add_paragraph(item, style='List Bullet')

# 4. 验证结果
doc.add_heading('4. 验证结果', level=1)

doc.add_heading('4.1 P0-P4 阶段性验证', level=2)
headers = ['Phase', '内容', '状态', '通过率/关键指标']
rows = [
    ['P0', '空间-语义统一', '✅ 通过', '分离比2.25, KNN 86.67%'],
    ['P1', '关键节点归并', '✅ 通过', '准确率100%'],
    ['P2', '双模式漫游', '✅ 通过', '空闲漫游14/14, 思考漫游75%/100%'],
    ['P3-FULL', '工作空间共振+晋升+降级', '✅ 通过', '共振60%, 双标准晋升, 退化验证'],
    ['P4', '端到端整合', '✅ 通过', '记忆100%, 漫游发现跨域联想'],
    ['P5', '三锚点+动态压缩栈', '✅ 通过', '压缩栈cascade验证, 唤起成本量化'],
]
table = doc.add_table(rows=1+len(rows), cols=4)
table.style = 'Table Grid'
for ci, h in enumerate(headers):
    table.rows[0].cells[ci].text = h
for ri, row in enumerate(rows):
    for ci, cell in enumerate(row):
        table.rows[ri+1].cells[ci].text = cell
doc.add_paragraph()

doc.add_heading('4.2 P5 动态压缩栈实验结果', level=2)
doc.add_paragraph('场景1-3：话题增加触发压缩级联')
for item in [
    '3话题（清晰区）：项目A(primary)、孩子学校(secondary)、跑步(tertiary) 全部full激活',
    '第4话题进入：项目A被压缩为summary(3线索)，激活度0.25，3秒后可唤起',
    '第5-6话题：孩子学校→跑步→项目A 依次压缩，最终3个旧话题均在压缩区(3线索)',
    '第7+话题：最早的compressed(summary) 降级为clue(单概念)，激活度<0.01时清除',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('场景4：唤起成本量化')
for item in [
    '项目A (hint="项目A")：[compressed] 2-3s → summary',
    '孩子学校 (hint="孩子")：[compressed] 2-3s → summary',
    '马拉松 (hint="马拉松")：[clear] instant → full',
    '账单缴费 (hint="账单")：[clear] instant → full',
    '父母：完全遗忘（不在清晰区也不在压缩区）',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('场景5：遗忘时间线（激活度衰减）')
for item in [
    'T+0：项目A被挤出压缩区，激活度0.25',
    'T+3h：激活度0.12，子节点从3→1',
    'T+9h：激活度<0.01 → 完全清除，代表性概念从栈中消失',
    '注：MSG中的项目A节点本身完好，只是想不起来了',
]:
    doc.add_paragraph(item, style='List Bullet')

# 5. 关键机制
doc.add_heading('5. 关键机制详解', level=1)

doc.add_heading('5.1 跨锚点桥梁检测', level=2)
doc.add_paragraph(
    '三锚点之间存在共同邻居节点时，系统主动发现跨域关联。'
    '实验发现：压力、孩子、跑步三个话题通过"时间"这个桥梁节点相连。'
    '系统输出建议："压力和孩子都涉及时间——你是否需要讨论时间管理策略？"'
)

doc.add_heading('5.2 压缩区与MSG的关系', level=2)
doc.add_paragraph(
    '压缩区是独立的映射空间，与MSG是读写分离的关系：'
)
for item in [
    '遗忘只发生在压缩区：激活度衰减、子节点删除、代表性概念清除',
    'MSG本身不受任何操作：节点、边、向量坐标从不删除',
    '唤起时：从压缩区找到代表性概念 → 映射回MSG取完整信息',
    '唤起成本：完全取决于压缩区状态，与MSG结构无关',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.3 平滑过渡与淘汰机制', level=2)
doc.add_paragraph(
    '老锚点不会突然消失，而是平滑过渡：'
    '老锚点激活度>0.2时保留为低权重tertiary；'
    '激活度持续衰减，低于阈值才被清除；'
    '清除后MSG不变，只是不再能通过压缩区唤起。'
)

# 6. 局限与未来
doc.add_heading('6. 局限性与未来方向', level=1)
for item in [
    '空间维度诅咒：128/384维在高语义复杂度下可能不足',
    '漫游质量依赖图拓扑：拓扑稀疏时联想能力受限',
    'LLM接口脆弱性：编码模型选择对语义区分度影响大',
    '未来：神经-符号深度融合、多智能体记忆、情感动机标记',
]:
    doc.add_paragraph(item, style='List Bullet')

# 结论
doc.add_heading('7. 结论', level=1)
doc.add_paragraph(
    'LDSG 通过三层架构（STM/MSG/WS）+ 三锚点动态压缩栈，'
    '在工程层面实现了类人记忆的关键特性：多任务并行、渐进遗忘与创造性联想。'
    'P0-P5 全部验证通过，核心机制（共振检测、双模式漫游、动态压缩栈、跨锚点桥梁发现）均已验证可行。'
    '系统从"单线程深挖"进化到"多线程广度+深度"，为构建具备类人记忆能力的认知系统提供了可行路径。'
)

# 保存
outpath = '论文_LDSG_v2.0_2026-04-03.docx'
doc.save(outpath)
print(f'论文已保存: {outpath}')
